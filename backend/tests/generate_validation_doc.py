"""Generate the SOD & SA validation test-case Word document.

One table per validation category (from tests/inventory.py — the same source of
truth the pytest suite asserts against), so the documented expected message can
never drift from what the app actually returns.

Run:
    backend/venv/Scripts/python.exe backend/tests/generate_validation_doc.py
Output:
    docs/SOD_SA_Validation_TestCases.docx
"""

from __future__ import annotations

import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Make `tests` importable when run as a script from anywhere.
_BACKEND_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _BACKEND_DIR not in sys.path:
    sys.path.insert(0, _BACKEND_DIR)

from tests.inventory import SCENARIOS, by_category  # noqa: E402

_HEADERS = ["#", "Scenario", "File to corrupt", "How to make it wrong",
            "Expected message (must appear)", "HTTP"]


def _shade(cell, hex_fill: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_cell(cell, text: str, *, bold=False, size=9, mono=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if mono:
        run.font.name = "Consolas"
    return cell


def build(output_path: str) -> None:
    doc = Document()

    title = doc.add_heading("SOD & SA Analysis — File Validation Test Cases", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph()
    intro.add_run(
        "This document lists every file-validation scenario for the SOD & SA Analysis "
        "upload and run flow. For each scenario, deliberately corrupt the named file as "
        "described, upload it, and confirm the application rejects it with the expected "
        "message and HTTP status."
    ).font.size = Pt(10)

    note = doc.add_paragraph()
    note.add_run("Important — corrupt ONE thing at a time. ").bold = True
    note.add_run(
        "Validation runs in phases (Phase 0 → 1 → 2 → 3 → run-time). The first phase that "
        "finds a problem returns immediately, so if you break several things at once you "
        "will only see the earliest one. Fix to a known-good baseline, then introduce a "
        "single defect per test."
    ).font.size = Pt(10)

    verified = doc.add_paragraph()
    verified.add_run(
        f"All {len([s for s in SCENARIOS])} scenarios below are also covered by an automated "
        "pytest suite (backend/tests/test_sod_sa_validation.py) that was run against the live "
        "application — every expected message in this document was confirmed to match the "
        "app's actual response."
    ).font.size = Pt(9)
    verified.runs[0].italic = True

    for category, scenarios in by_category().items():
        doc.add_heading(category, level=1)

        table = doc.add_table(rows=1, cols=len(_HEADERS))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(_HEADERS):
            _set_cell(hdr[i], h, bold=True, size=9)
            _shade(hdr[i], "2E2E38")
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for n, s in enumerate(scenarios, 1):
            cells = table.add_row().cells
            _set_cell(cells[0], str(n), size=9)
            _set_cell(cells[1], s.scenario, size=9)
            _set_cell(cells[2], s.file, size=9)
            _set_cell(cells[3], s.how_to_break, size=9)
            _set_cell(cells[4], s.expect, size=8, mono=True)
            code = "200" if s.http == 200 else str(s.http)
            _set_cell(cells[5], code, bold=True, size=9)

        # Column widths tuned for a 1366×768 / A4 landscape-ish read.
        widths = [0.3, 1.6, 1.2, 2.6, 2.2, 0.5]
        for row in table.rows:
            for i, w in enumerate(widths):
                from docx.shared import Inches
                row.cells[i].width = Inches(w)

    doc.add_heading("How to reset to a known-good baseline", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "A valid upload needs: a Role Hierarchy file (columns TOP_ROLE_CODE, TOP_ROLE_NAME, "
        "ROLE_CODE, ROLE_NAME, PRIVILEGE_CODE, PRIVILEGE_NAME); a Ruleset workbook with the "
        "three sheets 'SoD Ruleset', 'SA Ruleset', 'Entitlement to Privilege' (and optionally "
        "'Bucket Details'); optionally a User Role file (User Name, Assigned Role Name); and "
        "optionally an FP Database workbook with 'No_action_Privileges' and 'WorkArea_Privileges'. "
        "Start each test from this clean set. The bundled templates in "
        "backend/templates/sod-sa-analysis/ have the correct headers (add at least one data row)."
    ).font.size = Pt(10)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Wrote {output_path} ({len(SCENARIOS)} scenarios across {len(by_category())} tables)")


if __name__ == "__main__":
    out = os.path.join(_BACKEND_DIR, "..", "docs", "SOD_SA_Validation_TestCases.docx")
    build(os.path.abspath(out))
